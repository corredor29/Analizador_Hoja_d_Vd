import re
from docx import Document
from src.extractors.pdf_extractor import limpiar_texto
from src.parsers.contact_parser import parsear_contacto


def extraer_texto_docx(archivo) -> str:
    """
    Extrae todo el texto de un archivo Word (.docx).
    'archivo' puede ser una ruta (str) o un objeto de bytes (Streamlit UploadedFile).
    """
    documento = Document(archivo)
    texto = ""

    for parrafo in documento.paragraphs:
        if parrafo.text.strip():
            texto += parrafo.text + "\n"

    for tabla in documento.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                if celda.text.strip():
                    texto += celda.text + "\n"

    return limpiar_texto(texto)


def extraer_contacto_docx(archivo) -> dict:
    texto = extraer_texto_docx(archivo)
    return parsear_contacto(texto)